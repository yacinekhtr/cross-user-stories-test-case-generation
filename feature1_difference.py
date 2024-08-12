import os
import spacy
from docx import Document
from difflib import ndiff
from difflib import unified_diff
from difflib import SequenceMatcher

# Charger le modèle de langue en anglais de spaCy
nlp = spacy.load('en_core_web_sm')

def extract_text_from_docx(docx_path):
    doc = Document(docx_path)
    full_text = []

    # Extraire le texte des paragraphes
    for para in doc.paragraphs:
        full_text.append(para.text)
    
    # Extraire le texte des tableaux
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text for cell in row.cells]
            full_text.append(' '.join(row_text))
    
    return '\n'.join(full_text)

def get_user_stories_from_folder(folder_path):
    user_stories = {}
    for filename in os.listdir(folder_path):
        if filename.endswith(".docx"):
            file_path = os.path.join(folder_path, filename)
            text = extract_text_from_docx(file_path)
            user_stories[filename] = text
    return user_stories

def compare_user_stories(story1, story2):
    # Tokeniser les textes pour une comparaison plus précise
    doc1 = nlp(story1)
    doc2 = nlp(story2)
    tokens1 = [token.text for token in doc1 if not token.is_punct and not token.is_space]
    tokens2 = [token.text for token in doc2 if not token.is_punct and not token.is_space]
    
    matcher = SequenceMatcher(None, tokens1, tokens2)
    differences = []
    
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == 'replace':
            differences.append(f'Replace {tokens1[i1:i2]} with {tokens2[j1:j2]}')
        elif tag == 'delete':
            differences.append(f'Delete {tokens1[i1:i2]}')
        elif tag == 'insert':
            differences.append(f'Insert {tokens2[j1:j2]}')
    
    return differences

def analyze_differences_between_stories(stories):
    differences = {}
    story_keys = list(stories.keys())
    for i in range(len(story_keys)):
        for j in range(i + 1, len(story_keys)):
            story1_key = story_keys[i]
            story2_key = story_keys[j]
            differences[f"{story1_key} vs {story2_key}"] = compare_user_stories(stories[story1_key], stories[story2_key])
    return differences

# Chemin vers le dossier contenant les fichiers Word
folder_path = "C:\\Users\\YacineKHITER\\Documents\\User stories\\Paquet 1"

# Extraction des user stories
user_stories = get_user_stories_from_folder(folder_path)

# Vérification de l'extraction du texte
for filename, text in user_stories.items():
    print(f"Extracted text from {filename}:\n{text}\n{'='*80}\n")

# Analyse des différences entre les user stories
differences = analyze_differences_between_stories(user_stories)

# Affichage des différences
for comparison, diff in differences.items():
    print(f"Differences between {comparison}:")
    for d in diff:
        print(d)
    print("\n") 