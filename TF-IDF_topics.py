import os
import spacy
import numpy as np
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.cluster import KMeans
from sklearn.metrics import silhouette_score
from sklearn.metrics import confusion_matrix
from docx import Document

# Extraction du texte des fichiers .docx
def extract_text_from_docx(docx_path):
    doc = Document(docx_path)
    full_text = []

    # Extract text from paragraphs
    for para in doc.paragraphs:
        full_text.append(para.text)
    
    # Extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text for cell in row.cells]
            full_text.append(' '.join(row_text))
    
    return '\n'.join(full_text)

# Extraction des user stories à partir d'un dossier de fichiers .docx
def get_user_stories_from_folder(folder_path):
    user_stories = []
    for filename in os.listdir(folder_path):
        if filename.endswith(".docx"):
            file_path = os.path.join(folder_path, filename)
            text = extract_text_from_docx(file_path)
            user_stories.append(text)
    return user_stories

# Prétraitement des textes
def preprocess(text):
    result = []
    for token in nlp(text):
        if not token.is_stop and not token.is_punct and not token.is_space:
            result.append(token.lemma_.lower())
    return ' '.join(result)

# Création de la matrice TF-IDF
def create_tfidf_matrix(user_stories):
    vectorizer = TfidfVectorizer(max_df=0.95, min_df=2, stop_words='english')
    tfidf_matrix = vectorizer.fit_transform(user_stories)
    return tfidf_matrix, vectorizer

# Clustering des documents avec KMeans
def cluster_documents(tfidf_matrix, num_clusters=3, random_state=42):
    kmeans = KMeans(n_clusters=num_clusters, random_state=random_state)
    cluster_labels = kmeans.fit_predict(tfidf_matrix)
    return cluster_labels, kmeans

# Évaluation avec le score Silhouette
def evaluate_silhouette_score(tfidf_matrix, cluster_labels):
    silhouette_avg = silhouette_score(tfidf_matrix, cluster_labels)
    print(f"Silhouette Score: {silhouette_avg:.4f}")
    return silhouette_avg

# Calcul du Purity Score (nécessite des labels réels)
def purity_score(true_labels, cluster_labels):
    contingency_matrix = confusion_matrix(true_labels, cluster_labels)
    purity = np.sum(np.amax(contingency_matrix, axis=0)) / np.sum(contingency_matrix)
    print(f"Purity Score: {purity:.4f}")
    return purity

if __name__ == '__main__':
    # Chargement du modèle spacy
    nlp = spacy.load('en_core_web_sm')

    # Chemin vers le dossier contenant les fichiers .docx des user stories
    folder_path = "C:\\Users\\YacineKHITER\\Documents\\User stories\\Paquet 1"

    # Extraction des user stories
    user_stories = get_user_stories_from_folder(folder_path)

    # Prétraitement des user stories
    processed_stories = [preprocess(story) for story in user_stories]

    # Création de la matrice TF-IDF
    tfidf_matrix, vectorizer = create_tfidf_matrix(processed_stories)

    # Clustering des documents avec KMeans
    num_clusters = 3  # Vous pouvez ajuster ce paramètre en fonction du nombre de clusters attendu
    cluster_labels, kmeans_model = cluster_documents(tfidf_matrix, num_clusters=num_clusters, random_state=42)

    # Évaluation du modèle avec le score Silhouette
    silhouette_avg = evaluate_silhouette_score(tfidf_matrix, cluster_labels)

    # Si vous avez des étiquettes réelles (true_labels), vous pouvez calculer le Purity Score
    true_labels = []  # Remplacez cette liste par vos étiquettes réelles si disponibles
    if true_labels and len(true_labels) == len(cluster_labels):
        purity = purity_score(true_labels, cluster_labels)
    else:
        print("Aucun étiquetage réel fourni, skipping Purity Score calculation.")

    # Affichage des topics identifiés (mots-clés les plus importants par cluster)
    order_centroids = kmeans_model.cluster_centers_.argsort()[:, ::-1]
    terms = vectorizer.get_feature_names_out()
    for i in range(num_clusters):
        print(f"Cluster {i}:")
        for ind in order_centroids[i, :10]:  # Les 10 mots les plus fréquents
            print(f' {terms[ind]}')

    # Affichage des user stories dans chaque cluster
    for i, story in enumerate(user_stories):
        print(f'\nUser Story: {story[:100]}...')  # Affichage des 100 premiers caractères pour le contexte
        print(f'Cluster: {cluster_labels[i]}')
