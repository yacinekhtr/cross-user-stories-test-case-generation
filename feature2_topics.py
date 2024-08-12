import os
import spacy
from docx import Document
from gensim import corpora
from gensim.models.ldamodel import LdaModel
from gensim.parsing.preprocessing import preprocess_string, STOPWORDS
from difflib import ndiff, unified_diff
from sklearn.cluster import KMeans
import numpy as np
from gensim.models import CoherenceModel

def extract_text_from_docx(docx_path):
    doc = Document(docx_path)
    full_text = []

    # Extract text from paragraphs
    for para in doc.paragraphs:
        full_text.append(para.text)
    
    # Extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text for cell in row.cells]
            full_text.append(' '.join(row_text))
    
    return '\n'.join(full_text)

def get_user_stories_from_folder(folder_path):
    user_stories = []
    for filename in os.listdir(folder_path):
        if filename.endswith(".docx"):
            file_path = os.path.join(folder_path, filename)
            text = extract_text_from_docx(file_path)
            user_stories.append(text)
    return user_stories

def preprocess(text):
    result = []
    for token in nlp(text):
        if token.text.lower() not in STOPWORDS and not token.is_punct and not token.is_space:
            result.append(token.lemma_.lower())
    return result

def identify_topics(user_stories, num_topics=3, random_state=42):
    # Preprocess user stories
    processed_stories = [preprocess(story) for story in user_stories]
    
    # Create dictionary and the corpus
    dictionary = corpora.Dictionary(processed_stories)
    corpus = [dictionary.doc2bow(story) for story in processed_stories]
    
    # Use LDA with fixed random state
    lda_model = LdaModel(corpus, num_topics=num_topics, id2word=dictionary, passes=15, random_state=random_state)
    
    # Calculate coherence
    coherence_model_lda = CoherenceModel(model=lda_model, texts=processed_stories, dictionary=dictionary, coherence='c_v')
    coherence_lda = coherence_model_lda.get_coherence()
    
    return lda_model, dictionary, corpus, coherence_lda

def cluster_documents(lda_model, corpus, num_clusters=2, random_state=42):
    # Get document-topic distribution
    doc_topics = [lda_model.get_document_topics(bow) for bow in corpus]
    
    # Convert to a matrix
    doc_topic_matrix = np.zeros((len(doc_topics), lda_model.num_topics))
    for i, doc in enumerate(doc_topics):
        for topic_num, prob in doc:
            doc_topic_matrix[i, topic_num] = prob
    
    # Cluster the documents with fixed random state
    kmeans = KMeans(n_clusters=num_clusters, random_state=random_state).fit(doc_topic_matrix)
    return kmeans.labels_

def get_highest_coherence_topic(user_stories, lda_model, dictionary):
    processed_stories = [preprocess(story) for story in user_stories]
    coherence_values = []
    
    for i in range(lda_model.num_topics):
        topics = [[word for word, prob in lda_model.show_topic(i, topn=10)]]
        coherence_model = CoherenceModel(topics=topics, texts=processed_stories, dictionary=dictionary, coherence='c_v')
        coherence_values.append(coherence_model.get_coherence())
    
    highest_coherence_topic = np.argmax(coherence_values)
    return highest_coherence_topic

if __name__ == '__main__':
    # Chargement du modèle spacy
    nlp = spacy.load('en_core_web_sm')

    # Path to folder containing Word files
    folder_path = "C:\\Users\\YacineKHITER\\Documents\\User stories\\Paquet 1"

    # Extract user stories
    user_stories = get_user_stories_from_folder(folder_path)

    # Identify Topics
    lda_model, dictionary, corpus, coherence_lda = identify_topics(user_stories, random_state=42)

    # Cluster documents
    num_clusters = 1  # You can change this based on the number of expected clusters
    clusters = cluster_documents(lda_model, corpus, num_clusters=num_clusters, random_state=42)

    # Calculate topic coherence for each user story
    user_story_topics = []
    for story in user_stories:
        topics = lda_model.get_document_topics(dictionary.doc2bow(preprocess(story)))
        user_story_topics.append(topics)

    # Find the highest coherence topic
    highest_coherence_topic_value = 0
    highest_coherence_topic = None
    for i, story_topics in enumerate(user_story_topics):
        for topic_num, topic_value in story_topics:
            if topic_value > highest_coherence_topic_value:
                highest_coherence_topic_value = topic_value
                highest_coherence_topic = topic_num

    # Display the topics
    for idx, topic in lda_model.print_topics(-1):
        print(f'Topic {idx}: {topic}')

    # Display Topics and Cluster from a specific user story
    for i, story in enumerate(user_stories):
        print(f'\nUser Story: {story[:100]}...')  # Display the first 100 characters for context
        print(f'Assigned Common Topic: {highest_coherence_topic}')
        print(f'Cluster: {clusters[i]}')

    # Display coherence
    print(f'Coherence of the LDA model: {coherence_lda:.4f}')

    # Display detailed topics for the user story with the highest topic probability
    for i, story in enumerate(user_stories):
        topics = lda_model.get_document_topics(dictionary.doc2bow(preprocess(story)))
        highest_topic = max(topics, key=lambda x: x[1])
        print(f'\nUser Story: {story[:100]}...')  # Display the first 100 characters for context
        print('Topics:', topics)
        print('Highest Topic:', highest_topic)
